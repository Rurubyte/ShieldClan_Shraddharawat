from io import BytesIO

# pyrefly: ignore [missing-import]
import pytest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.core.errors import AppException
from app.services.onboarding.attachment_builder import AttachmentBuilder
from app.services.onboarding.email_template_builder import EmailTemplateBuilder, EmailTemplateContext
from app.services.onboarding.interview_session_builder import InterviewSessionBuilder
from app.services.resume.candidate_extractor import CandidateExtractor
from app.services.resume.demo_assets import generate_demo_resume_pdf
from app.services.resume.docx_parser import extract_text_from_docx
from app.services.resume.parser import ResumeParserService
from app.services.resume.pdf_parser import extract_text_from_pdf


def _build_sample_pdf() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.drawString(72, 720, "Jane Doe")
    pdf.drawString(72, 700, "jane.doe@example.com")
    pdf.drawString(72, 680, "9876543210")
    pdf.drawString(72, 660, "https://linkedin.com/in/janedoe")
    pdf.drawString(72, 640, "https://github.com/janedoe")
    pdf.drawString(72, 600, "Skills")
    pdf.drawString(72, 580, "Python, FastAPI, PostgreSQL")
    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def _build_sample_docx() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane.doe@example.com")
    document.add_paragraph("9876543210")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI")
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def test_pdf_parsing_extracts_text():
    text = extract_text_from_pdf(_build_sample_pdf())
    assert "Jane Doe" in text
    assert "jane.doe@example.com" in text


def test_docx_parsing_extracts_text():
    text = extract_text_from_docx(_build_sample_docx())
    assert "Jane Doe" in text
    assert "jane.doe@example.com" in text


def test_contact_extraction_from_resume_text():
    parsed = CandidateExtractor().extract(
        "Jane Doe\njane.doe@example.com\n9876543210\nhttps://github.com/janedoe"
    )
    assert parsed.name == "Jane Doe"
    assert parsed.email == "jane.doe@example.com"
    assert parsed.phone == "9876543210"
    assert parsed.github_url == "https://github.com/janedoe"


def test_resume_parser_rejects_unsupported_format():
    with pytest.raises(AppException) as exc:
        ResumeParserService().parse_bytes(b"hello", "resume.txt")
    assert exc.value.status_code == 400


def test_email_template_renders_html_and_subject():
    builder = EmailTemplateBuilder()
    subject, html, plain = builder.build(
        EmailTemplateContext(
            candidate_name="Jane Doe",
            position_name="Backend Engineer",
            resume_score=91,
            shortlist_reasons=["ML Systems"],
            interview_topics=["RAG"],
            interview_instructions=["Be on time"],
            interview_url="http://localhost/interview",
            link_expiry="2026-01-01T10:00:00+00:00",
            company_name="ICD",
            recruiter_contact="recruiter@example.com",
            skills=["Python", "FastAPI"],
            interview_mode="Virtual Video Call",
            interview_duration="60 mins",
        )
    )
    assert "Interview Invitation" in subject
    assert "Backend Engineer" in subject
    assert "Jane Doe" in html
    assert "RAG" in html
    assert "Python" in html
    assert "FastAPI" in html
    assert "Virtual Video Call" in html
    assert "60 mins" in html
    assert "Jane Doe" in plain
    assert "Python" not in plain  # plain text fallback does not necessarily render badges


def test_email_template_optional_fields_hidden():
    builder = EmailTemplateBuilder()
    # Build context without skills, position, mode, duration, shortlist_reasons, or interview_topics
    subject, html, plain = builder.build(
        EmailTemplateContext(
            candidate_name="Jane Doe",
            position_name="",
            resume_score=85,
            shortlist_reasons=[],
            interview_topics=[],
            interview_instructions=["Be on time"],
            interview_url="http://localhost/interview",
            link_expiry="2026-01-01T10:00:00+00:00",
            company_name="ICD",
            recruiter_contact="recruiter@example.com",
        )
    )
    assert "Interview Invitation" in subject
    assert "Backend Engineer" not in subject
    assert "Jane Doe" in html
    assert "Skills Matched" not in html
    assert "Why You Were Shortlisted" not in html
    assert "Interview Focus Topics" not in html
    assert "Interview Mode" not in html
    assert "Interview Duration" not in html
    assert "Position" not in html



def test_interview_url_generation():
    url = InterviewSessionBuilder.build_interview_url(
        base_url="http://localhost:8000",
        session_uuid="11111111-1111-1111-1111-111111111111",
        raw_token="abc-token",
    )
    assert url.endswith("token=abc-token")
    assert "session=11111111-1111-1111-1111-111111111111" in url


def test_jd_attachment_generates_placeholder_when_missing(tmp_path):
    builder = AttachmentBuilder(jd_pdf_path=str(tmp_path / "missing.pdf"), generated_dir=str(tmp_path / "generated"))
    path = builder.resolve_jd_attachment(position_name="Engineer", company_name="ICD")
    assert path.exists()
    assert path.suffix == ".pdf"


def test_demo_resume_pdf_generation():
    content = generate_demo_resume_pdf()
    text = extract_text_from_pdf(content)
    assert "rahul@gmail.com" in text
